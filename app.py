# Version 3.10 Claude
"""FastAPI service that converts Markdown to Word and returns a download URL.
Supports headings, bullet lists, numbered lists, bold/italic, links, inline code, 
code blocks, and GitHub‑style tables with borders + bold header row. 
Files saved under /static/downloads/ ."""

from fastapi import FastAPI, Body
from fastapi.staticfiles import StaticFiles
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.table import _Cell
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from uuid import uuid4
import os
import re

app = FastAPI()
BASE_DIR = "generated"
SUB_DIR = "downloads"
OUTPUT_DIR = os.path.join(BASE_DIR, SUB_DIR)

os.makedirs(OUTPUT_DIR, exist_ok=True)
app.mount("/static", StaticFiles(directory=BASE_DIR), name="static")

# Regex patterns
_hdr1 = re.compile(r"^# (.+)")
_hdr2 = re.compile(r"^## (.+)")
_hdr3 = re.compile(r"^### (.+)")
_bullet = re.compile(r"^[-*] (.+)")
_numbered = re.compile(r"^(\d+)\. (.+)")
_nested_bullet = re.compile(r"^  [-*] (.+)")
_nested_numbered = re.compile(r"^  (\d+)\. (.+)")
_table_row = re.compile(r"^\|(.+)\|$")
_table_divider = re.compile(r"^[\|\s:-]+$")
_code_block = re.compile(r"^```")
_bold = re.compile(r"\*\*(.+?)\*\*")
_italic = re.compile(r"\*(.+?)\*")
_inline_code = re.compile(r"`(.+?)`")
_link = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_html_strong = re.compile(r"<strong>(.+?)</strong>", re.IGNORECASE)
_invalid_fname = re.compile(r"[^A-Za-z0-9_.-]")

# Helper functions

def _set_cell_border(cell: _Cell):
    tcPr = cell._tc.get_or_add_tcPr()
    for edge in ("top", "left", "bottom", "right"):
        ln = OxmlElement(f"w:{edge}")
        ln.set(qn("w:val"), "single")
        ln.set(qn("w:sz"), "4")
        ln.set(qn("w:space"), "0")
        ln.set(qn("w:color"), "auto")
        tcPr.append(ln)

def _safe_filename(name: str | None) -> str:
    if not name:
        return f"{uuid4()}.docx"
    name = _invalid_fname.sub("_", name).strip("._") or str(uuid4())
    if not name.lower().endswith(".docx"):
        name += ".docx"
    return name

def _inline_formats(par, text: str):
    """Process inline formatting including bold, italic, inline code, and links"""
    pos = 0
    while pos < len(text):
        # Find all possible matches
        matches = []
        
        for pattern, format_type in [
            (_bold, 'bold'),
            (_italic, 'italic'), 
            (_inline_code, 'code'),
            (_link, 'link')
        ]:
            match = pattern.search(text, pos)
            if match:
                matches.append((match.start(), match, format_type))
        
        if not matches:
            # No more formatting, add remaining text
            par.add_run(text[pos:])
            break
        
        # Sort by position to handle earliest match first
        matches.sort(key=lambda x: x[0])
        earliest_pos, match, format_type = matches[0]
        
        # Add text before the match
        if earliest_pos > pos:
            par.add_run(text[pos:earliest_pos])
        
        # Add formatted text
        if format_type == 'link':
            # For links, create hyperlink
            link_text = match.group(1)
            link_url = match.group(2)
            hyperlink = par.add_run(link_text)
            # Add hyperlink relationship
            r_id = par.part.relate_to(link_url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
            hyperlink_elem = OxmlElement('w:hyperlink')
            hyperlink_elem.set(qn('r:id'), r_id)
            new_run = OxmlElement('w:r')
            new_run.append(hyperlink._element)
            hyperlink_elem.append(new_run)
            par._element.append(hyperlink_elem)
        else:
            run = par.add_run(match.group(1))
            if format_type == 'bold':
                run.bold = True
            elif format_type == 'italic':
                run.italic = True
            elif format_type == 'code':
                # Plain text for inline code as requested
                pass
        
        pos = match.end()

def _is_valid_table_structure(lines, start_idx):
    """Check if we have a valid table structure starting at start_idx"""
    if start_idx >= len(lines):
        return False
    
    # Must have header row
    if not _table_row.match(lines[start_idx]):
        return False
    
    # Must have divider row
    if start_idx + 1 >= len(lines) or not _table_divider.match(lines[start_idx + 1]):
        return False
    
    # Must have at least one data row
    if start_idx + 2 >= len(lines) or not _table_row.match(lines[start_idx + 2]):
        return False
    
    return True

def md_to_docx(md: str, path: str):
    doc = Document()
    
    # Ensure required styles exist
    style_names = [s.name for s in doc.styles]
    if "List Bullet" not in style_names:
        doc.styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH).base_style = doc.styles["Normal"]
    if "List Number" not in style_names:
        doc.styles.add_style("List Number", WD_STYLE_TYPE.PARAGRAPH).base_style = doc.styles["Normal"]
    if "List Bullet 2" not in style_names:
        bullet2_style = doc.styles.add_style("List Bullet 2", WD_STYLE_TYPE.PARAGRAPH)
        bullet2_style.base_style = doc.styles["Normal"]
    if "List Number 2" not in style_names:
        number2_style = doc.styles.add_style("List Number 2", WD_STYLE_TYPE.PARAGRAPH)
        number2_style.base_style = doc.styles["Normal"]

    lines = md.splitlines()
    i = 0
    in_code_block = False
    code_content = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if _code_block.match(line):
            if in_code_block:
                # End of code block - add as plain text
                if code_content:
                    p = doc.add_paragraph()
                    for code_line in code_content:
                        p.add_run(code_line + "\n")
                code_content = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
        
        # Handle tables - but only if structure is valid
        if _table_row.match(line) and _is_valid_table_structure(lines, i):
            header = [c.strip() for c in line.strip("|").split("|")]
            
            # Skip divider row
            i += 2
            
            # Collect data rows
            rows = []
            while i < len(lines) and _table_row.match(lines[i]):
                row_data = [c.strip() for c in lines[i].strip("|").split("|")]
                # Ensure row has same number of columns as header
                while len(row_data) < len(header):
                    row_data.append("")
                rows.append(row_data[:len(header)])  # Truncate if too many columns
                i += 1
            
            # Create table
            tbl = doc.add_table(rows=len(rows)+1, cols=len(header))
            
            # Add header row
            for c, txt in enumerate(header):
                cell = tbl.rows[0].cells[c]
                cell.text = txt
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                _set_cell_border(cell)
            
            # Add data rows
            for r, row in enumerate(rows, start=1):
                for c, txt in enumerate(row):
                    cell = tbl.rows[r].cells[c]
                    cell.text = txt
                    _set_cell_border(cell)
            
            continue
        
        # Handle headings
        if m := _hdr1.match(line):
            doc.add_heading(m.group(1), level=1)
        elif m := _hdr2.match(line):
            doc.add_heading(m.group(1), level=2)
        elif m := _hdr3.match(line):
            doc.add_heading(m.group(1), level=3)
        # Handle nested lists
        elif m := _nested_bullet.match(line):
            p = doc.add_paragraph()
            p.style = "List Bullet 2"
            _inline_formats(p, m.group(1))
        elif m := _nested_numbered.match(line):
            p = doc.add_paragraph()
            p.style = "List Number 2"  
            _inline_formats(p, m.group(2))
        # Handle top-level lists
        elif m := _bullet.match(line):
            p = doc.add_paragraph()
            p.style = "List Bullet"
            _inline_formats(p, m.group(1))
        elif m := _numbered.match(line):
            p = doc.add_paragraph()
            p.style = "List Number"
            _inline_formats(p, m.group(2))
        else:
            # Regular paragraph
            plain = _html_strong.sub(r"**\1**", line)
            if plain.strip():  # Only add non-empty paragraphs
                p = doc.add_paragraph()
                _inline_formats(p, plain)
        
        i += 1
    
    # Handle any remaining code block content
    if in_code_block and code_content:
        p = doc.add_paragraph()
        for code_line in code_content:
            p.add_run(code_line + "\n")
    
    doc.save(path)

@app.post("/docx")
def make_docx(payload: dict = Body(...)):
    md = payload["markdown"]
    fname = _safe_filename(payload.get("filename"))
    full = os.path.join(OUTPUT_DIR, fname)
    md_to_docx(md, full)
    host = os.environ.get("RENDER_EXTERNAL_HOSTNAME", "localhost")
    return {"download_url": f"https://{host}/static/{SUB_DIR}/{fname}"}
